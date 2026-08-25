#!/usr/bin/env python3
"""Produce a lightweight 《经济研究》 submission-format audit for a DOCX."""

import argparse
import json
import re
from pathlib import Path

from docx import Document


PATTERNS = {
    "has_content_abstract_label": r"内容摘要：",
    "has_chinese_abstract": r"(内容提要|内容摘要|摘要：)",
    "has_keywords": r"关键词：",
    "has_summary": r"\bSummary\b|英文梗概",
    "has_key_words": r"Key Words\s*:",
    "has_jel": r"JEL\s+Classification|JEL分类|JEL 分类",
    "has_references": r"参考文献",
    "has_tables": r"表\s*\d+",
    "has_figures": r"图\s*\d+",
    "has_equation_number": r"[（(]\s*\d+\s*[）)]",
}

IDENTITY_PATTERNS = {
    "email": r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}",
    "phone_like": r"(?:\+?86[-\s]?)?1[3-9]\d{9}",
    "fund_note": r"(基金项目|资助项目|项目编号|国家自然科学基金|教育部|社科基金)",
    "acknowledgement": r"(致谢|感谢|匿名审稿人)",
}

METADATA_FIELDS = ("author", "last_modified_by", "title", "subject", "comments", "keywords", "category")


def collect_text(doc):
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(p.text for p in cell.paragraphs)
    return "\n".join(parts)


def collect_contexts(doc, pattern, limit=8):
    contexts = []
    for index, paragraph in enumerate(doc.paragraphs, start=1):
        text = paragraph.text.strip()
        if text and re.search(pattern, text, re.IGNORECASE):
            contexts.append({"location": f"paragraph {index}", "text": text[:240]})
            if len(contexts) >= limit:
                return contexts
    for table_index, table in enumerate(doc.tables, start=1):
        for row_index, row in enumerate(table.rows, start=1):
            for cell_index, cell in enumerate(row.cells, start=1):
                cell_text = "\n".join(p.text.strip() for p in cell.paragraphs if p.text.strip())
                if cell_text and re.search(pattern, cell_text, re.IGNORECASE):
                    contexts.append({
                        "location": f"table {table_index} row {row_index} cell {cell_index}",
                        "text": cell_text[:240],
                    })
                    if len(contexts) >= limit:
                        return contexts
    return contexts


def collect_metadata(doc):
    props = doc.core_properties
    metadata = {}
    for field in METADATA_FIELDS:
        value = getattr(props, field, None)
        if value:
            metadata[field] = str(value)
    return metadata


def audit(path):
    doc = Document(path)
    text = collect_text(doc)
    report = {
        "file": str(path),
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(doc.tables),
        "detected_items": {},
        "possible_identity_leaks": {},
        "identity_contexts": {},
        "metadata": collect_metadata(doc),
        "warnings": [],
    }

    for name, pattern in PATTERNS.items():
        report["detected_items"][name] = bool(re.search(pattern, text, re.IGNORECASE))

    for name, pattern in IDENTITY_PATTERNS.items():
        matches = re.findall(pattern, text, re.IGNORECASE)
        report["possible_identity_leaks"][name] = len(matches)
        if matches:
            report["identity_contexts"][name] = collect_contexts(doc, pattern)

    if not report["detected_items"]["has_chinese_abstract"]:
        report["warnings"].append("Missing Chinese content abstract marker.")
    if not report["detected_items"]["has_content_abstract_label"]:
        report["warnings"].append("Expected screenshot-style label 内容摘要： was not found.")
    if not report["detected_items"]["has_keywords"]:
        report["warnings"].append("Missing Chinese keywords marker.")
    if not report["detected_items"]["has_summary"]:
        report["warnings"].append("Missing English Summary marker.")
    if not report["detected_items"]["has_key_words"]:
        report["warnings"].append("Missing Key Words marker.")
    if not report["detected_items"]["has_jel"]:
        report["warnings"].append("Missing JEL Classification marker.")
    if not report["detected_items"]["has_references"]:
        report["warnings"].append("Missing references heading.")
    if any(report["possible_identity_leaks"].values()):
        report["warnings"].append("Possible identity information found; check anonymous-review requirements.")
    if report["metadata"]:
        report["warnings"].append("Document metadata is not empty; check anonymous-review requirements.")

    return report


def main():
    parser = argparse.ArgumentParser(description="Audit a DOCX for common ERJ submission-format items.")
    parser.add_argument("input", help="Input DOCX path")
    parser.add_argument("--json", action="store_true", help="Print JSON instead of plain text")
    args = parser.parse_args()

    report = audit(Path(args.input))
    if args.json:
        print(json.dumps(report, ensure_ascii=False, indent=2))
        return

    print(f"File: {report['file']}")
    print(f"Paragraphs: {report['paragraph_count']}")
    print(f"Tables: {report['table_count']}")
    print("Detected:")
    for key, value in report["detected_items"].items():
        print(f"  - {key}: {value}")
    print("Possible identity leaks:")
    for key, value in report["possible_identity_leaks"].items():
        print(f"  - {key}: {value}")
    if report["identity_contexts"]:
        print("Identity contexts:")
        for key, contexts in report["identity_contexts"].items():
            print(f"  - {key}:")
            for context in contexts:
                print(f"    {context['location']}: {context['text']}")
    if report["metadata"]:
        print("Metadata:")
        for key, value in report["metadata"].items():
            print(f"  - {key}: {value}")
    if report["warnings"]:
        print("Warnings:")
        for warning in report["warnings"]:
            print(f"  - {warning}")


if __name__ == "__main__":
    main()
