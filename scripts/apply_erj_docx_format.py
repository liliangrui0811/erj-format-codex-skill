#!/usr/bin/env python3
"""Apply a conservative 《经济研究》-style DOCX format pass.

This helper is intentionally conservative. It normalizes common typography,
heading patterns, captions, and table borders, but it does not replace careful
manual inspection against a current journal template.
"""

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def ensure_rpr(run):
    run._element.get_or_add_rPr()
    if run._element.rPr.rFonts is None:
        run._element.rPr.append(OxmlElement("w:rFonts"))


DEFAULT_EAST_FONT = "宋体"
DEFAULT_WEST_FONT = "Times New Roman"
DEFAULT_ABSTRACT_FONT = "仿宋"
DEFAULT_TABLE_FONT = "仿宋"
DEFAULT_FIGURE_FONT = "黑体"
DEFAULT_REFERENCE_FONT = "宋体"


def set_run_font(run, east=DEFAULT_EAST_FONT, west=DEFAULT_WEST_FONT, size=10.5, bold=None, italic=None):
    ensure_rpr(run)
    run.font.name = west
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east)
    run._element.rPr.rFonts.set(qn("w:ascii"), west)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), west)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_para_font(paragraph, east=DEFAULT_EAST_FONT, west=DEFAULT_WEST_FONT, size=10.5, bold=None):
    if not paragraph.runs:
        paragraph.add_run("")
    for run in paragraph.runs:
        set_run_font(run, east=east, west=west, size=size, bold=bold)


def set_cell_border(cell, **kwargs):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge, edge_data in kwargs.items():
        element = borders.find(qn("w:" + edge))
        if element is None:
            element = OxmlElement("w:" + edge)
            borders.append(element)
        for key, value in edge_data.items():
            element.set(qn("w:" + key), str(value))


def apply_three_line_table(table):
    nil = {"val": "nil", "sz": "0", "color": "auto"}
    line = {"val": "single", "sz": "8", "color": "000000", "space": "0"}
    for row in table.rows:
        row_pr = row._tr.get_or_add_trPr()
        if row_pr.find(qn("w:cantSplit")) is None:
            row_pr.append(OxmlElement("w:cantSplit"))
        for cell in row.cells:
            set_cell_border(cell, top=nil, left=nil, bottom=nil, right=nil, insideH=nil, insideV=nil)
    if table.rows:
        for cell in table.rows[0].cells:
            set_cell_border(cell, top=line, bottom=line)
        for cell in table.rows[-1].cells:
            set_cell_border(cell, bottom=line)


def apply_paragraph_format(paragraph, first_line=True, line_spacing=1.0):
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = line_spacing
    paragraph.paragraph_format.first_line_indent = Cm(0.74) if first_line else Cm(0)


def add_two_space_prefix(paragraph):
    text = paragraph.text
    if not text or text.startswith(("  ", "　　")):
        return
    first_run = paragraph.runs[0] if paragraph.runs else paragraph.add_run("")
    first_run.text = "  " + first_run.text


def is_probable_author_line(index, stripped):
    if index not in {1, 2}:
        return False
    if stripped.startswith(("——", "--", "内容摘要", "内容提要", "摘要：", "关键词", "English Summary")):
        return False
    if "匿名审稿版" in stripped:
        return False
    if len(stripped) > 40:
        return False
    return not re.search(r"[，。；：,.!?！？]", stripped)


def classify_paragraph(index, text):
    stripped = text.strip()
    if not stripped:
        return "blank"
    if index == 0:
        return "title"
    if index == 1 and stripped.startswith(("——", "--")):
        return "subtitle"
    if is_probable_author_line(index, stripped):
        return "author"
    if "匿名审稿版" in stripped:
        return "anonymous_note"
    if stripped.startswith(("内容提要", "内容摘要", "摘要：", "关键词")):
        return "cn_frontmatter"
    if stripped.startswith(("Key Words", "JEL Classification")):
        return "frontmatter"
    if stripped.startswith("注"):
        return "note"
    if stripped in {"参考文献", "References"}:
        return "references_heading"
    if re.match(r"^[一二三四五六七八九十]+、", stripped):
        return "heading1"
    if re.match(r"^（[一二三四五六七八九十]+）", stripped):
        return "heading2"
    if re.match(r"^\d+[.．]\s*", stripped):
        return "heading3"
    if re.match(r"^表\s*\d+", stripped):
        return "table_caption"
    if re.match(r"^图\s*\d+", stripped):
        return "figure_caption"
    return "body"


def format_docx(
    input_path,
    output_path,
    body_size=10.5,
    cjk_font=DEFAULT_EAST_FONT,
    western_font=DEFAULT_WEST_FONT,
    abstract_font=DEFAULT_ABSTRACT_FONT,
    table_font=DEFAULT_TABLE_FONT,
    figure_font=DEFAULT_FIGURE_FONT,
    reference_font=DEFAULT_REFERENCE_FONT,
):
    doc = Document(input_path)

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    for style_name in ("Normal", "No Spacing"):
        if style_name not in doc.styles:
            continue
        style = doc.styles[style_name]
        style.font.name = western_font
        style._element.rPr.rFonts.set(qn("w:eastAsia"), cjk_font)
        style._element.rPr.rFonts.set(qn("w:ascii"), western_font)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), western_font)
        style.font.size = Pt(body_size)
        style.paragraph_format.first_line_indent = Cm(0.74)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)

    in_references = False
    for index, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text
        kind = classify_paragraph(index, text)
        if kind == "blank":
            continue

        if "Normal" in doc.styles:
            paragraph.style = doc.styles["Normal"]

        if kind == "title":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            apply_paragraph_format(paragraph, first_line=False)
            set_para_font(paragraph, east=cjk_font, west=western_font, size=15, bold=True)
        elif kind == "subtitle":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            apply_paragraph_format(paragraph, first_line=False)
            set_para_font(paragraph, east=cjk_font, west=western_font, size=12, bold=False)
        elif kind == "author":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            apply_paragraph_format(paragraph, first_line=False)
            set_para_font(paragraph, east=cjk_font, west=western_font, size=12, bold=False)
        elif kind == "anonymous_note":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            apply_paragraph_format(paragraph, first_line=False)
            set_para_font(paragraph, east=cjk_font, west=western_font, size=9, bold=False)
        elif kind == "cn_frontmatter":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            apply_paragraph_format(paragraph, first_line=True)
            set_para_font(paragraph, east=abstract_font, west=western_font, size=10.5, bold=False)
        elif kind == "frontmatter":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            apply_paragraph_format(paragraph, first_line=True)
            set_para_font(paragraph, east=cjk_font, west=western_font, size=10.5, bold=False)
        elif kind == "references_heading":
            in_references = True
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            apply_paragraph_format(paragraph, first_line=False)
            set_para_font(paragraph, east=cjk_font, west=western_font, size=10.5, bold=True)
        elif kind == "heading1":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            apply_paragraph_format(paragraph, first_line=False)
            set_para_font(paragraph, east=cjk_font, west=western_font, size=14, bold=False)
        elif kind in {"heading2", "heading3"}:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            apply_paragraph_format(paragraph, first_line=False)
            set_para_font(paragraph, east=cjk_font, west=western_font, size=body_size, bold=False)
        elif kind == "table_caption":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            apply_paragraph_format(paragraph, first_line=False)
            set_para_font(paragraph, east=cjk_font, west=western_font, size=9, bold=True)
        elif kind == "figure_caption":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            apply_paragraph_format(paragraph, first_line=False)
            set_para_font(paragraph, east=figure_font, west=western_font, size=9, bold=True)
        elif kind == "note":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            apply_paragraph_format(paragraph, first_line=False)
            set_para_font(paragraph, east=reference_font, west=western_font, size=8, bold=False)
        elif in_references:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            apply_paragraph_format(paragraph, first_line=False)
            add_two_space_prefix(paragraph)
            set_para_font(paragraph, east=reference_font, west=western_font, size=7.5, bold=False)
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            apply_paragraph_format(paragraph, first_line=True)
            set_para_font(paragraph, east=cjk_font, west=western_font, size=body_size, bold=False)

    for table in doc.tables:
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.autofit = True
        apply_three_line_table(table)
        for row_index, row in enumerate(table.rows):
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    apply_paragraph_format(paragraph, first_line=False)
                    set_para_font(paragraph, east=table_font, west=western_font, size=9, bold=(row_index == 0))

    doc.core_properties.author = ""
    doc.core_properties.last_modified_by = ""
    doc.core_properties.title = ""
    doc.core_properties.subject = ""
    doc.core_properties.comments = ""
    doc.core_properties.keywords = ""
    doc.core_properties.category = ""

    doc.save(output_path)


def main():
    parser = argparse.ArgumentParser(description="Apply conservative ERJ DOCX formatting.")
    parser.add_argument("input", help="Input DOCX path")
    parser.add_argument("--out", required=True, help="Output DOCX path")
    parser.add_argument("--body-size", type=float, default=10.5, help="Body font size in points")
    parser.add_argument("--cjk-font", default=DEFAULT_EAST_FONT, help="Chinese font name")
    parser.add_argument("--western-font", default=DEFAULT_WEST_FONT, help="Western font name")
    parser.add_argument("--abstract-font", default=DEFAULT_ABSTRACT_FONT, help="Chinese font name for abstract and keywords")
    parser.add_argument("--table-font", default=DEFAULT_TABLE_FONT, help="Chinese font name for tables")
    parser.add_argument("--figure-font", default=DEFAULT_FIGURE_FONT, help="Chinese font name for figure captions")
    parser.add_argument("--reference-font", default=DEFAULT_REFERENCE_FONT, help="Chinese font name for references and notes")
    args = parser.parse_args()

    input_path = Path(args.input)
    output_path = Path(args.out)
    if input_path.resolve() == output_path.resolve():
        raise SystemExit("Refusing to overwrite the input file. Choose a separate --out path.")
    format_docx(
        input_path,
        output_path,
        body_size=args.body_size,
        cjk_font=args.cjk_font,
        western_font=args.western_font,
        abstract_font=args.abstract_font,
        table_font=args.table_font,
        figure_font=args.figure_font,
        reference_font=args.reference_font,
    )


if __name__ == "__main__":
    main()
